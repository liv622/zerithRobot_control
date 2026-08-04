#!/usr/bin/env python3
from pathlib import Path
import sys

sys.path.insert(0, "/tmp/e1pro_docx")

from lxml import html
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor
from PIL import Image


ROOT = Path(__file__).resolve().parents[1]
SOURCE = Path(sys.argv[1]).resolve() if len(sys.argv) > 1 else ROOT / "deliverables" / "vacuum_gripper_research_report.html"
OUTPUT = Path(sys.argv[2]).resolve() if len(sys.argv) > 2 else ROOT / "deliverables" / "商超零售_小型吸夹复合末端执行器调研报告_V1.0.docx"
IS_V2 = "solution_report_v2" in SOURCE.name


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_border(cell, color="B7C6CD", size="4"):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:color"), color)


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def add_hyperlink(paragraph, text, url):
    part = paragraph.part
    rel_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), rel_id)
    new_run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "087E9B")
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    r_pr.append(color)
    r_pr.append(underline)
    new_run.append(r_pr)
    text_node = OxmlElement("w:t")
    text_node.text = text
    new_run.append(text_node)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


def set_run_font(run, size=None, bold=None, color=None, mono=False):
    run.font.name = "DejaVu Sans Mono" if mono else "Noto Sans CJK SC"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), run.font.name)
    if size:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)


def add_inline(paragraph, node, inherited_bold=False):
    if node.text:
        run = paragraph.add_run(node.text)
        set_run_font(run, bold=inherited_bold or node.tag in ("b", "strong"))
        if node.tag == "sub":
            run.font.subscript = True
    for child in node:
        bold = inherited_bold or node.tag in ("b", "strong")
        if child.tag == "a" and child.get("href"):
            add_hyperlink(paragraph, "".join(child.itertext()), child.get("href"))
        elif child.tag == "br":
            paragraph.add_run().add_break()
        else:
            add_inline(paragraph, child, bold)
        if child.tail:
            run = paragraph.add_run(child.tail)
            set_run_font(run, bold=inherited_bold)


def add_rich_paragraph(doc, element, style=None, alignment=None):
    paragraph = doc.add_paragraph(style=style)
    if alignment:
        paragraph.alignment = alignment
    add_inline(paragraph, element)
    return paragraph


def add_callout(doc, element, kind):
    colors = {
        "callout": ("EAF6F8", "1496A8"),
        "warning": ("FFF4E8", "E08A32"),
        "danger": ("FCEDEE", "C94B54"),
        "formula": ("F3F5F7", "D3DCE1"),
        "diagram": ("F8FAFB", "C9D6DC"),
    }
    fill, border = colors[kind]
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    set_cell_border(cell, border, "8" if kind in ("callout", "warning", "danger") else "4")
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    paragraph = cell.paragraphs[0]
    if kind in ("diagram", "formula"):
        run = paragraph.add_run(element.text_content().strip())
        set_run_font(run, size=8.5 if kind == "diagram" else 9, mono=True)
    else:
        add_inline(paragraph, element)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def add_list(doc, element, ordered=False):
    items = element.xpath("./li")
    reference_list = bool(items) and all((li.get("id") or "").startswith("R") for li in items)
    for li in items:
        p = doc.add_paragraph(style=None if reference_list else ("List Number" if ordered else "List Bullet"))
        if reference_list:
            p.paragraph_format.left_indent = Cm(0.45)
            p.paragraph_format.first_line_indent = Cm(-0.45)
        add_inline(p, li)


def add_html_table(doc, element):
    rows = element.xpath("./tr|./thead/tr|./tbody/tr")
    if not rows:
        return
    col_count = max(len(row.xpath("./th|./td")) for row in rows)
    table = doc.add_table(rows=len(rows), cols=col_count)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    for i, row in enumerate(rows):
        cells = row.xpath("./th|./td")
        for j, source_cell in enumerate(cells):
            target_cell = table.cell(i, j)
            target_cell.text = ""
            target_cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            p = target_cell.paragraphs[0]
            add_inline(p, source_cell)
            for run in p.runs:
                set_run_font(run, size=8.5, bold=(i == 0 or source_cell.tag == "th"))
                if i == 0 or source_cell.tag == "th":
                    run.font.color.rgb = RGBColor(255, 255, 255)
            set_cell_shading(target_cell, "176B87" if i == 0 or source_cell.tag == "th" else ("F4F8FA" if i % 2 else "FFFFFF"))
            set_cell_border(target_cell)
        if i == 0:
            set_repeat_table_header(table.rows[i])
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def add_figure(doc, element):
    image_nodes = element.xpath("./img")
    if not image_nodes:
        return
    source = SOURCE.parent / image_nodes[0].get("src")
    with Image.open(source) as image:
        ratio = image.width / image.height
    max_width = 15.6
    max_height = 11.8
    if max_width / ratio <= max_height:
        picture = doc.add_picture(str(source), width=Cm(max_width))
    else:
        picture = doc.add_picture(str(source), height=Cm(max_height))
    picture_paragraph = doc.paragraphs[-1]
    picture_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    picture_paragraph.paragraph_format.keep_with_next = True
    captions = element.xpath("./figcaption")
    if captions:
        caption = doc.add_paragraph(style="Caption")
        caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
        caption.paragraph_format.keep_together = True
        add_inline(caption, captions[0])
        for run in caption.runs:
            set_run_font(run, size=8.5, color=(88, 102, 114))


def add_cover(doc, section):
    for element in section:
        if element.tag == "div" and "tag" in (element.get("class") or ""):
            continue
        if element.tag == "h1":
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(105)
            p.paragraph_format.space_after = Pt(18)
            text = element.text_content().replace("\n", " ")
            run = p.add_run(text)
            set_run_font(run, size=24, bold=True, color=(23, 59, 87))
        elif element.tag == "div" and "subtitle" in (element.get("class") or ""):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(element.text_content())
            set_run_font(run, size=15, color=(71, 105, 122))
        elif element.tag == "p":
            p = add_rich_paragraph(doc, element)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif element.tag == "div" and "meta" in (element.get("class") or ""):
            for child in element.xpath("./p"):
                p = add_rich_paragraph(doc, child)
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in p.runs:
                    set_run_font(run, size=9, color=(92, 105, 115))
    doc.add_page_break()


def setup_document():
    doc = Document()
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(1.8)
    section.bottom_margin = Cm(1.6)
    section.left_margin = Cm(1.8)
    section.right_margin = Cm(1.8)

    normal = doc.styles["Normal"]
    normal.font.name = "Noto Sans CJK SC"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Noto Sans CJK SC")
    normal.font.size = Pt(10.5)
    normal.paragraph_format.line_spacing = 1.35
    normal.paragraph_format.space_after = Pt(4)

    for style_name, size, color in (
        ("Title", 24, (23, 59, 87)),
        ("Heading 1", 16, (23, 107, 135)),
        ("Heading 2", 12.5, (36, 86, 110)),
        ("Heading 3", 11, (54, 95, 115)),
    ):
        style = doc.styles[style_name]
        style.font.name = "Noto Sans CJK SC"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Noto Sans CJK SC")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor(*color)
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.space_before = Pt(11)
        style.paragraph_format.space_after = Pt(5)

    for style_name in ("List Bullet", "List Number"):
        style = doc.styles[style_name]
        style.font.name = "Noto Sans CJK SC"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Noto Sans CJK SC")
        style.font.size = Pt(10)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    version = "V2.1" if IS_V2 else "V1.0"
    run = footer.add_run(f"商超零售小型吸夹复合末端执行器方案｜{version}｜2026-07-31")
    set_run_font(run, size=8, color=(105, 119, 128))
    return doc


def convert():
    tree = html.parse(str(SOURCE))
    body = tree.getroot().xpath("//body")[0]
    doc = setup_document()

    for element in body:
        tag = element.tag.lower() if isinstance(element.tag, str) else ""
        cls = element.get("class") or ""
        if tag == "section" and "cover" in cls:
            add_cover(doc, element)
        elif tag == "h2":
            doc.add_heading(element.text_content().strip(), level=1)
        elif tag == "h3":
            doc.add_heading(element.text_content().strip(), level=2)
        elif tag == "h4":
            doc.add_heading(element.text_content().strip(), level=3)
        elif tag == "p":
            p = add_rich_paragraph(doc, element)
            if "small" in cls or "footer-note" in cls:
                for run in p.runs:
                    set_run_font(run, size=8.5, color=(88, 102, 114))
        elif tag == "ul":
            add_list(doc, element, ordered=False)
        elif tag == "ol":
            add_list(doc, element, ordered=True)
        elif tag == "table":
            add_html_table(doc, element)
        elif tag == "figure":
            add_figure(doc, element)
        elif tag == "div" and "pagebreak" in cls:
            doc.add_page_break()
        elif tag == "div":
            for kind in ("callout", "warning", "danger", "formula", "diagram"):
                if kind in cls:
                    add_callout(doc, element, kind)
                    break

    core = doc.core_properties
    core.title = "商超零售场景小型吸夹复合末端执行器可行方案与定量选型" if IS_V2 else "商超零售场景小型吸夹复合末端执行器调研与设计边界"
    core.subject = "吸盘模块、吸附动作、夹爪集成、定量能力、风险边界、规划控制与验证" if IS_V2 else "真空吸盘边界、A4纸单张吸取、夹爪集成、规划控制、标准与验证"
    core.author = "Codex"
    core.keywords = "真空吸盘, 夹爪, 末端执行器, 商超零售, 软包装, 异形商品, 抓取规划"
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    convert()
